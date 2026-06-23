import os
import pandas as pd
from docx import Document
from docx.shared import Inches
from datetime import datetime
from pathlib import Path

# ==========================
# RUTAS DEL PROYECTO
# ==========================

CARPETA_PCAPS = r"T:\Universidad\Tesis\datasets\Modbus Dataset"

CARPETA_RESUMEN = r"T:\Universidad\Tesis\datasets\resumen"

CARPETA_GRAFICOS = r"T:\Universidad\Tesis\datasets\graficos"

RUTA_ESTADISTICAS = r"T:\Universidad\Tesis\datasets\estadisticas.json"

ARCHIVO_SALIDA = r"T:\Universidad\Tesis\datasets\codios de reporte\Reporte_Forense_Modbus.docx"

# ==========================
# INFORMACION DEL ENTORNO
# ==========================

print("\n=== INFORMACION DEL ENTORNO ===\n")

sistema_operativo = input(
    "Sistema Operativo: "
)

version_python = input(
    "Version Python: "
)

version_zeek = input(
    "Version Zeek: "
)

version_timesketch = input(
    "Version Timesketch: "
)

autor = input("Autor: ")

# ==========================================
# FUNCIONES
# ==========================================

def contar_pcaps(ruta):

    total_archivos = 0
    total_bytes = 0

    for root, dirs, files in os.walk(ruta):

        for f in files:

            if f.endswith(".pcap"):

                total_archivos += 1
                total_bytes += os.path.getsize(
                    os.path.join(root, f)
                )

    return total_archivos, total_bytes


def agregar_tabla(doc, titulo, archivo_csv):

    if not os.path.exists(archivo_csv):
        return

    doc.add_heading(titulo, level=2)

    df = pd.read_csv(archivo_csv)

    filas = min(len(df), 15)

    tabla = doc.add_table(
        rows=filas + 1,
        cols=len(df.columns)
    )

    tabla.style = "Table Grid"

    for j, col in enumerate(df.columns):
        tabla.cell(0, j).text = str(col)

    for i in range(filas):
        for j in range(len(df.columns)):
            tabla.cell(
                i + 1,
                j
            ).text = str(df.iloc[i, j])

    doc.add_paragraph("")


def agregar_imagen(doc, titulo, imagen):

    if not os.path.exists(imagen):
        return

    doc.add_heading(titulo, level=2)

    doc.add_picture(
        imagen,
        width=Inches(6)
    )

    doc.add_paragraph("")


# ==========================================
# DOCUMENTO
# ==========================================

doc = Document()

doc.add_heading(
    "Reporte Forense Modbus/TCP",
    level=1
)

# ==========================================
# ENTORNO
# ==========================================

doc.add_heading(
    "1. Entorno Experimental",
    level=1
)

doc.add_paragraph(f"Autor: {autor}")
doc.add_paragraph(f"Fecha: {datetime.now()}")
doc.add_paragraph(f"Sistema Operativo: {sistema_operativo}")
doc.add_paragraph(f"Python: {version_python}")
doc.add_paragraph(f"Zeek: {version_zeek}")
doc.add_paragraph(f"Timesketch: {version_timesketch}")

# ==========================================
# DATASET
# ==========================================

doc.add_heading(
    "2. Estadisticas del Dataset",
    level=1
)

pcaps, total_bytes = contar_pcaps(CARPETA_PCAPS)

doc.add_paragraph(
    f"Cantidad de PCAPs: {pcaps}"
)

doc.add_paragraph(
    f"Tamaño total KB: {total_bytes/1024:.2f}"
)

doc.add_paragraph(
    f"Tamaño total MB: {total_bytes/(1024**2):.2f}"
)

doc.add_paragraph(
    f"Tamaño total GB: {total_bytes/(1024**3):.2f}"
)

# ==========================================
# GRAFICOS
# ==========================================

GRAF = CARPETA_GRAFICOS
RES = CARPETA_RESUMEN


agregar_imagen(
    doc,
    "3. Errores Detectados por Analyzer",
    os.path.join(GRAF, "analyzer.png")
)

# ==========================================
# ORIGEN
# ==========================================

doc.add_heading(
    "4. IPs de Origen",
    level=1
)

agregar_tabla(
    doc,
    "Origen Benigno",
    os.path.join(
        RES,
        "origen_benigno.csv"
    )
)

agregar_tabla(
    doc,
    "Origen Ataque",
    os.path.join(
        RES,
        "origen_ataque.csv"
    )
)

agregar_imagen(
    doc,
    "IPs de Origen mas Comunes",
    os.path.join(
        GRAF,
        "ip_origen.png"
    )
)

# ==========================================
# DESTINO
# ==========================================

doc.add_heading(
    "5. IPs de Destino",
    level=1
)

agregar_tabla(
    doc,
    "Destino Benigno",
    os.path.join(
        RES,
        "destino_benigno.csv"
    )
)

agregar_tabla(
    doc,
    "Destino Ataque",
    os.path.join(
        RES,
        "destino_ataque.csv"
    )
)

agregar_imagen(
    doc,
    "IPs de Destino mas Comunes",
    os.path.join(
        GRAF,
        "ip_destino.png"
    )
)

# ==========================================
# MODBUS
# ==========================================

doc.add_heading(
    "6. Funciones Modbus",
    level=1
)

agregar_tabla(
    doc,
    "Funciones Benignas",
    os.path.join(
        RES,
        "modbus_benigno.csv"
    )
)

agregar_tabla(
    doc,
    "Funciones Ataque",
    os.path.join(
        RES,
        "modbus_ataque.csv"
    )
)

agregar_imagen(
    doc,
    "Funciones Modbus",
    os.path.join(
        GRAF,
        "modbus_funciones.png"
    )
)

agregar_imagen(
    doc,
    "Solicitudes y Respuestas",
    os.path.join(
        GRAF,
        "modbus_req_resp.png"
    )
)

# ==========================================
# WEIRD
# ==========================================

doc.add_heading(
    "7. Eventos Weird",
    level=1
)

agregar_tabla(
    doc,
    "Weird Benigno",
    os.path.join(
        RES,
        "weird_benigno.csv"
    )
)

agregar_tabla(
    doc,
    "Weird Ataque",
    os.path.join(
        RES,
        "weird_ataque.csv"
    )
)

agregar_imagen(
    doc,
    "Eventos Weird",
    os.path.join(
        GRAF,
        "weird.png"
    )
)

# ==========================================
# TIMELINE
# ==========================================

doc.add_heading(
    "8. Analisis Temporal",
    level=1
)

agregar_imagen(
    doc,
    "Linea Temporal",
    os.path.join(
        GRAF,
        "timeline.png"
    )
)

# ==========================================
# CONCLUSIONES
# ==========================================

doc.add_heading(
    "9. Conclusiones",
    level=1
)

doc.add_paragraph(
"""
Se observaron diferencias significativas entre el tráfico benigno y el tráfico asociado a escenarios de ataque.

Las direcciones IP de origen muestran la aparición de la IP 185.175.0.7 como un actor relevante durante los escenarios maliciosos.

Las funciones WRITE_SINGLE_COIL y WRITE_SINGLE_REGISTER presentan incrementos importantes durante los ataques, evidenciando intentos de modificación del proceso industrial.

Los eventos Weird identificados indican anomalías de comunicación consistentes con actividades de exploración, manipulación y reutilización de conexiones.

El análisis temporal evidencia picos de actividad que coinciden con las fases de ataque observadas en el dataset.
"""
)

salida = ARCHIVO_SALIDA
print("\nIMAGENES ENCONTRADAS:")

for archivo in os.listdir(CARPETA_GRAFICOS):
    print(archivo)
doc.save(salida)

print()
print("================================")
print("REPORTE GENERADO")
print("================================")
print(salida)
print("================================")